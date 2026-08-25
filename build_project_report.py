from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Clio_RAG_Project_Report_draft.docx"


def set_font(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text))
    return p


def add_prompt(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.08
    set_font(p.add_run(label + "\n"), size=10, bold=True, color="434343")
    set_font(p.add_run(text), size=10, color="434343")
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "DADCE0")
    borders.append(left)
    p._p.get_or_add_pPr().append(borders)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size, before, after, color in (
    ("Heading 1", 20, 20, 6, "000000"),
    ("Heading 2", 16, 18, 6, "000000"),
    ("Heading 3", 14, 16, 4, "434343"),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
set_font(title.add_run("Clio: A Source-Grounded History Tutor"), size=26)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_font(subtitle.add_run("Project report | LangChain, LangGraph, Gemini, and Streamlit"), size=11, color="555555")
meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(20)
set_font(meta.add_run("Prepared for project submission | August 24, 2026"), size=10, color="555555")

doc.add_heading("1. Project overview", level=1)
add_body(doc, "Clio is a retrieval-augmented generation (RAG) application that turns a private collection of history documents into a question-answering tutor. A user uploads a PDF, DOCX, TXT, or Markdown file, asks a question, and receives an answer grounded only in passages retrieved from those documents. Answers retain filename and page or document-location metadata so the user can inspect the evidence behind each citation.")
add_body(doc, "The application was built to make the retrieval process understandable and testable. It implements two chunking strategies over the same source collection, retrieves the six most similar chunks for the same query, and can display both results side by side. This makes it possible to examine whether topic-aware semantic boundaries improve retrieval over a conventional fixed-size baseline.")

doc.add_heading("Core capabilities", level=2)
for item in (
    "Local parsing of PDF, DOCX, TXT, and Markdown files, with Gemini OCR as a fallback for image-only PDFs.",
    "A fixed-size chunking baseline and an embedding-based semantic chunking strategy.",
    "Dense vector retrieval with Gemini embeddings and cosine-style similarity ranking.",
    "A LangGraph retrieve-then-generate workflow with source-only answer instructions.",
    "Side-by-side answers and retrieved evidence for fixed and semantic strategies.",
    "Persistent SQLite storage for extracted documents, chunks, and embeddings across restarts.",
    "A labelled evaluation harness that reports Hit@k and mean reciprocal rank (MRR).",
):
    add_bullet(doc, item)

doc.add_heading("2. System workflow", level=1)
for step in (
    "Ingest documents. The Streamlit interface accepts one or more supported files. PDF text is extracted page by page; scanned PDFs can be transcribed through Gemini OCR.",
    "Preserve provenance. Every extracted section becomes a LangChain Document with filename and page or document-location metadata.",
    "Create two indexes. The same source text is processed by fixed chunking and semantic chunking, producing two directly comparable collections.",
    "Embed and persist. Gemini converts chunks into vectors. Extracted documents, chunk sets, and embeddings are stored in rag_history.db and reused after a restart.",
    "Retrieve evidence. The user question is embedded and compared with each strategy's vector store. The six closest chunks are selected independently for each strategy.",
    "Generate grounded answers. LangGraph executes a retrieve node followed by a generate node. Gemini receives only the selected passages and must cite factual claims as [Source N].",
    "Compare. In Compare both mode, the app presents fixed and semantic answers and evidence in parallel.",
):
    add_number(doc, step)

doc.add_heading("Retrieval strategies", level=2)
add_body(doc, "Fixed chunks. Text is divided into approximately 1,200-character passages with 180 characters of overlap. The splitter prefers paragraph and sentence boundaries. This method is inexpensive, predictable, and serves as the baseline, but a topic may be divided at an arbitrary length boundary.", bold_lead="Fixed chunks.")
add_body(doc, "Semantic chunks. Text is separated around sentence-level changes in meaning. Sentence embeddings are compared, and a new chunk begins when the semantic distance is in the largest 15% of observed changes (the 85th-percentile breakpoint). A minimum chunk size of 120 characters prevents very small fragments. This approach can keep related ideas together, but initial indexing requires more embedding work.", bold_lead="Semantic chunks.")
add_body(doc, "Search method. Both indexes use the same dense retrieval pattern. The question and chunks are represented as vectors, ranked by vector similarity, and truncated to the top six. Therefore, the comparison isolates chunk construction rather than changing the search algorithm.", bold_lead="Search method.")

doc.add_heading("3. Datasets and source material", level=1)
add_body(doc, "No fixed public benchmark dataset is bundled with the repository. The operational dataset is the private document collection supplied by the user at runtime. During development, the workflow was exercised with history-oriented source material, including a user-uploaded PDF of approximately 3.5 MB, plus short synthetic text passages in automated tests. The application does not search the public web when answering questions.")
doc.add_heading("Supported data formats", level=2)
for item in (
    "PDF: page-aware extraction through pypdf; Gemini OCR fallback when pages contain no extractable text.",
    "DOCX: paragraph extraction through python-docx.",
    "TXT and Markdown: UTF-8 decoding with replacement for malformed characters.",
    "Evaluation JSON: document filenames plus labelled queries and distinctive relevant-text excerpts.",
):
    add_bullet(doc, item)
add_body(doc, "The included evaluation.example.json is a schema example, not an evaluation result. A meaningful quality comparison requires a representative set of questions and human-labelled relevant passages from the actual source collection.")

doc.add_heading("4. Prompts and agent instructions", level=1)
add_body(doc, "The most important instruction is the grounded-generation prompt used by the LangGraph generation node. It constrains the model to retrieved evidence and specifies a readable response structure.")
add_prompt(doc, "Grounded-answer system instruction", "You are Clio, a careful history tutor. Answer only from the supplied sources and cite every factual claim as [Source N]. If evidence is insufficient, say what is missing. Return readable Markdown, never JSON. Use these sections: 'Short answer' (2-3 sentences), 'Key points' (concise bullets), and 'Evidence digest' (one brief bullet per source actually used).")
add_body(doc, "The model receives the source passages and the question in a single human message. Each passage is labelled with its source number, filename, and location before the text is appended.")
add_prompt(doc, "Runtime message template", "SOURCES\n[Source 1: filename, location]\n<retrieved passage>\n\n...\n\nQUESTION\n<user question>")
add_body(doc, "For scanned PDFs, the OCR request focuses on transcription rather than interpretation so source fidelity is preserved.")
add_prompt(doc, "OCR instruction", "Transcribe all readable text in this PDF exactly. Preserve paragraph order. Return every page, including pages with no text. Do not summarize, explain, or add facts.")
doc.add_heading("Embedding instructions", level=2)
for item in (
    "Document passages use retrieval-document embeddings.",
    "Questions use retrieval-query embeddings.",
    "The original inspectable pipeline requests 768-dimensional Gemini embeddings.",
    "The LangChain pipeline uses GoogleGenerativeAIEmbeddings and caches vectors by embedding model, input type, and content hash.",
):
    add_bullet(doc, item)

doc.add_heading("5. Iterations tried", level=1)
iterations = (
    ("Inspectable baseline", "The first version used custom Python functions for boundary-aware fixed chunking, Gemini embeddings, direct cosine similarity, SQLite embedding storage, and a source-only Gemini prompt. This established a transparent baseline without framework abstraction."),
    ("LangChain and LangGraph refactor", "The second version represented sources as LangChain Documents, created separate vector stores, and introduced a LangGraph StateGraph with explicit retrieve and generate nodes."),
    ("Semantic chunking", "An embedding-based SemanticChunker was added with an 85th-percentile breakpoint and a 120-character minimum. Fixed chunks were retained so changes could be compared rather than assumed to be better."),
    ("Retrieval evaluation", "A command-line evaluation harness was added. It runs identical labelled queries through both retrievers and computes Hit@k and MRR."),
    ("Readable response formatting", "Gemini occasionally returned structured content blocks or stringified Python-like lists containing provider metadata. A normalizer now extracts only the text, converts real JSON into Markdown, and hides internal signature metadata."),
    ("Side-by-side comparison", "The Streamlit interface was expanded with Compare both mode, presenting answers and retrieved passages in two columns."),
    ("Persistent libraries", "SQLite persistence was extended to source documents, both chunk sets, and embeddings. Saved libraries can be reloaded after a server restart without repeating semantic chunking or document embedding API calls."),
    ("Environment and network troubleshooting", "A project-local virtual environment was created for stable imports. The app was also restarted with normal outbound network access after a restricted process caused a Gemini connection-refused error."),
)
for title_text, description in iterations:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(title_text + ": "), bold=True)
    set_font(p.add_run(description))

doc.add_heading("6. Evaluation approach", level=1)
add_body(doc, "Retrieval quality cannot be inferred reliably from whether an answer sounds fluent. The project therefore evaluates retrieval against human-labelled evidence. For each query, relevant_text contains one or more distinctive excerpts expected to appear in a relevant retrieved chunk.")
doc.add_heading("Metrics", level=2)
add_body(doc, "Hit@k measures whether at least one labelled relevant passage appears within the top k results. It answers: did the retriever find useful evidence at all?", bold_lead="Hit@k")
add_body(doc, "Mean reciprocal rank (MRR) rewards systems that place the first relevant result near the top. A relevant result at rank 1 contributes 1.0; at rank 2 it contributes 0.5; at rank 3 it contributes approximately 0.33.", bold_lead="Mean reciprocal rank (MRR)")
add_body(doc, "The evaluation also reports chunk counts because semantic chunking can change index granularity. Both strategies receive the same query set and cutoff, making the comparison controlled. No final claim that one strategy wins is included because the repository does not yet contain a completed, representative labelled dataset.")

doc.add_heading("7. Learnings and observations", level=1)
for item in (
    "Chunking and retrieval are separate decisions. The two strategies differ in passage boundaries, while both use the same vector-similarity retrieval method.",
    "Semantic chunking is not automatically superior. It may improve topic coherence, but it costs more during first-time indexing and can produce variable chunk sizes.",
    "A fixed baseline is essential. Without it, retrieval improvements cannot be attributed specifically to semantic boundaries.",
    "Retrieval quality needs labels. Side-by-side answers are useful for qualitative inspection, but Hit@k and MRR require human judgement about which passages are relevant.",
    "Provider output must be normalized. Model SDKs may return text blocks rather than plain strings, and blindly converting them with str() can expose internal metadata.",
    "Persistence materially improves usability. Caching only final vectors is not enough for semantic chunking; caching extracted documents and chunk sets prevents repeated preprocessing after restarts.",
    "Source metadata is part of the product, not an afterthought. Filenames and page locations let users audit citations and reduce the risk of accepting unsupported answers.",
    "Operational environment matters. Correct application code can still fail when the selected Python interpreter lacks dependencies or when the process cannot reach the model API.",
):
    add_bullet(doc, item)

doc.add_heading("8. Limitations and next steps", level=1)
for item in (
    "Create a larger labelled history-query set spanning factual, causal, comparative, and timeline questions.",
    "Add precision@k or nDCG when multiple relevance grades are available.",
    "Evaluate reranking, maximal marginal relevance, or hybrid keyword-vector retrieval as additional controlled variants.",
    "Replace the experimental semantic chunker if its upstream package is retired; current tests already surface its deprecation warning.",
    "Add document-library deletion and storage-size controls for long-term multi-user use.",
    "Consider answer-level faithfulness evaluation separately from retrieval evaluation.",
):
    add_bullet(doc, item)

doc.add_heading("9. Implementation summary", level=1)
add_body(doc, "The finished project is a source-grounded history tutor that combines a transparent baseline with a framework-based RAG workflow. LangChain manages document abstractions, semantic splitting, embeddings, and vector stores; LangGraph makes the retrieve-to-generate sequence explicit; Gemini supplies OCR, embeddings, and grounded generation; Streamlit provides the upload, comparison, and evidence-review interface; and SQLite supplies restart-safe persistence.")
add_body(doc, "The strongest outcome of the workflow is not simply that the app can answer questions. It also exposes the evidence, compares two retrieval designs under the same conditions, and provides a path to quantitative evaluation instead of relying only on subjective impressions.")

doc.add_heading("Appendix: key project files", level=1)
for item in (
    "app.py - Streamlit interface, saved-library loading, comparison layout, and answer rendering.",
    "langchain_rag.py - LangChain chunking/vector stores, LangGraph workflow, persistent cache, response normalization, and evaluation metrics.",
    "rag.py - inspectable baseline utilities, parsing, OCR, direct embeddings, cosine similarity, and source-grounded generation.",
    "evaluate_retrieval.py - command-line benchmark runner.",
    "evaluation.example.json - example labelled-query schema.",
    "tests/ - offline tests for chunking, retrieval metrics, response formatting, OCR metadata, and cache reuse.",
):
    add_bullet(doc, item)

doc.save(OUTPUT)
print(OUTPUT)
